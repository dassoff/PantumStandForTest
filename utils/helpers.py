"""Вспомогательные утилиты."""

import os
import shutil
import tempfile
import time
from pathlib import Path
from typing import Any, Dict, Optional


def format_bytes(size: int) -> str:
    """Форматирование размера в байтах в человекочитаемый вид."""
    for unit in ["B", "KB", "MB", "GB", "TB"]:
        if size < 1024.0:
            return f"{size:.2f} {unit}"
        size /= 1024.0
    return f"{size:.2f} PB"


def format_duration(ms: int) -> str:
    """Форматирование длительности в миллисекундах в человекочитаемый вид."""
    if ms < 1000:
        return f"{ms}ms"
    elif ms < 60000:
        return f"{ms / 1000:.2f}s"
    else:
        minutes = ms // 60000
        seconds = (ms % 60000) / 1000
        return f"{minutes}m {seconds:.1f}s"


def ensure_dir(path: str) -> Path:
    """Создание директории если она не существует."""
    p = Path(path)
    p.mkdir(parents=True, exist_ok=True)
    return p


def cleanup_temp(temp_dir: Optional[str] = None, older_than_hours: int = 1) -> int:
    """Очистка временных файлов."""
    if temp_dir is None:
        temp_dir = tempfile.gettempdir()

    temp_path = Path(temp_dir)
    if not temp_path.exists():
        return 0

    cutoff = time.time() - (older_than_hours * 3600)
    removed = 0

    for file in temp_path.glob("print_*"):
        try:
            if file.stat().st_mtime < cutoff:
                file.unlink()
                removed += 1
        except (OSError, PermissionError):
            pass

    return removed


def generate_test_pdf(
    output_path: str,
    pages: int = 1,
    text: str = "Test Print Document",
) -> str:
    """Генерация тестового PDF файла."""
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.pdfgen import canvas
    except ImportError:
        # Fallback: создаем простой PDF вручную
        return _generate_simple_pdf(output_path, pages, text)

    c = canvas.Canvas(output_path, pagesize=A4)
    width, height = A4

    for page in range(pages):
        c.setFont("Helvetica", 12)
        c.drawString(100, height - 100, f"{text} - Page {page + 1}")
        c.drawString(100, height - 120, f"Generated at: {time.strftime('%Y-%m-%d %H:%M:%S')}")

        # Добавляем немного текста для объема
        for i in range(20):
            c.drawString(100, height - 160 - (i * 20), f"Line {i + 1}: Test content for printing")

        if page < pages - 1:
            c.showPage()

    c.save()
    return output_path


def _generate_simple_pdf(output_path: str, pages: int, text: str) -> str:
    """Простая генерация PDF без внешних зависимостей."""
    # Минимальный валидный PDF с одной страницей
    pdf_content = f"""%PDF-1.4
1 0 obj
<< /Type /Catalog /Pages 2 0 R >>
endobj
2 0 obj
<< /Type /Pages /Kids [{"".join(f"{i+3} 0 R " for i in range(pages))}] /Count {pages} >>
endobj
"""

    for i in range(pages):
        pdf_content += f"""{i+3} 0 obj
<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] /Contents {i+pages+3} 0 R /Resources << /Font << /F1 {i+pages*2+3} 0 R >> >> >>
endobj
{i+pages+3} 0 obj
<< /Length 100 >>
stream
BT /F1 12 Tf 100 700 Td ({text} - Page {i+1}) Tj ET
endstream
endobj
{i+pages*2+3} 0 obj
<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>
endobj
"""

    pdf_content += f"""{pages*3+3} 0 obj
<< /Type /XRef /W [1 2 2] /Size {pages*3+4} /Index [0 {pages*3+4}] >>
stream
"""
    pdf_content += "\x00" * 20
    pdf_content += """
endstream
endobj
trailer
<< /Size """ + str(pages * 3 + 4) + """ /Root 1 0 R >>
%%EOF
"""

    with open(output_path, "wb") as f:
        f.write(pdf_content.encode())

    return output_path


def generate_test_docx(
    output_path: str,
    pages: int = 1,
    text: str = "Test DOCX Document",
) -> str:
    """Генерация тестового DOCX файла."""
    try:
        from docx import Document
        from docx.shared import Pt
    except ImportError:
        # Fallback: создаем простой DOCX
        return _generate_simple_docx(output_path, pages, text)

    doc = Document()

    # Заголовок
    heading = doc.add_heading(text, 0)

    # Метаинформация
    doc.add_paragraph(f"Generated at: {time.strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph(f"Pages: {pages}")

    # Контент для каждой "страницы"
    for page in range(pages):
        doc.add_page_break()
        doc.add_heading(f"Page {page + 1}", level=1)

        for i in range(30):
            para = doc.add_paragraph(f"Paragraph {i + 1}: ")
            para.add_run(text).bold = True
            para.add_run(" - This is test content for document printing validation.")

    doc.save(output_path)
    return output_path


def _generate_simple_docx(output_path: str, pages: int, text: str) -> str:
    """Простая генерация DOCX без внешних зависимостей."""
    # Минимальный DOCX файл (ZIP с XML)
    import zipfile

    # Создаем временную директорию
    temp_dir = tempfile.mkdtemp()

    try:
        # _rels/.rels
        rels_content = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
<Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="word/document.xml"/>
</Relationships>
"""
        with open(os.path.join(temp_dir, "_rels", ".rels"), "w") as f:
            f.write(rels_content)

        # word/document.xml
        doc_content = f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
<w:body>
<w:p><w:r><w:t>{text}</w:t></w:r></w:p>
<w:p><w:r><w:t>Generated at: {time.strftime('%Y-%m-%d %H:%M:%S')}</w:t></w:r></w:p>
"""
        for i in range(pages * 10):
            doc_content += f'<w:p><w:r><w:t>Line {i + 1}: Test content</w:t></w:r></w:p>'

        doc_content += """</w:body></w:document>"""

        word_dir = os.path.join(temp_dir, "word")
        os.makedirs(word_dir, exist_ok=True)

        with open(os.path.join(word_dir, "document.xml"), "w") as f:
            f.write(doc_content)

        # [Content_Types].xml
        types_content = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">
<Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>
<Default Extension="xml" ContentType="application/xml"/>
<Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>
</Types>
"""
        with open(os.path.join(temp_dir, "[Content_Types].xml"), "w") as f:
            f.write(types_content)

        # Создаем ZIP (DOCX)
        with zipfile.ZipFile(output_path, "w", zipfile.ZIP_DEFLATED) as zipf:
            for root, dirs, files in os.walk(temp_dir):
                for file in files:
                    file_path = os.path.join(root, file)
                    arcname = os.path.relpath(file_path, temp_dir)
                    zipf.write(file_path, arcname)

        return output_path

    finally:
        shutil.rmtree(temp_dir, ignore_errors=True)


def get_file_info(path: str) -> Dict[str, Any]:
    """Получение информации о файле."""
    p = Path(path)
    if not p.exists():
        return {"error": "File not found"}

    return {
        "name": p.name,
        "path": str(p.absolute()),
        "size": p.stat().st_size,
        "size_formatted": format_bytes(p.stat().st_size),
        "created": p.stat().st_ctime,
        "modified": p.stat().st_mtime,
        "extension": p.suffix.lower(),
    }


async def async_copy_file(src: str, dst: str) -> str:
    """Асинхронное копирование файла."""
    import aiofiles
    import aiofiles.os

    async with aiofiles.open(src, "rb") as f_src:
        content = await f_src.read()

    async with aiofiles.open(dst, "wb") as f_dst:
        await f_dst.write(content)

    return dst
