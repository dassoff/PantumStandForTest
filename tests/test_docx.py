"""Тесты DOCX печати."""

import os
import tempfile
from typing import Optional

import pytest

from core.printer import FastPrinter, PrintStatus
from utils.helpers import generate_test_docx


class TestDOCXPrinting:
    """Тесты печати DOCX файлов."""

    @pytest.fixture
    def printer(self, printer_ip: Optional[str] = None) -> FastPrinter:
        """Создание экземпляра принтера."""
        return FastPrinter(printer_ip=printer_ip)

    @pytest.fixture
    def sample_docx(self) -> str:
        """Создание тестового DOCX файла."""
        fd, path = tempfile.mkstemp(suffix=".docx")
        os.close(fd)
        generate_test_docx(path, pages=1, text="Test DOCX - Simple Text")
        yield path
        os.unlink(path)

    @pytest.fixture
    def docx_with_tables(self) -> str:
        """Создание DOCX с таблицами."""
        try:
            from docx import Document
        except ImportError:
            # Fallback на простую генерацию
            fd, path = tempfile.mkstemp(suffix=".docx")
            os.close(fd)
            generate_test_docx(path, pages=1, text="Table Test")
            return path

        doc = Document()
        doc.add_heading("DOCX with Tables", 0)

        # Добавляем таблицу
        table = doc.add_table(rows=3, cols=3)
        table.style = 'Table Grid'

        for i, row in enumerate(table.rows):
            for j, cell in enumerate(row.cells):
                cell.text = f"Cell {i+1},{j+1}"

        fd, path = tempfile.mkstemp(suffix=".docx")
        os.close(fd)
        doc.save(path)
        yield path
        os.unlink(path)

    @pytest.fixture
    def docx_with_images(self) -> str:
        """Создание DOCX с изображениями."""
        try:
            from docx import Document
            from docx.shared import Inches
            from PIL import Image
            import io
        except ImportError:
            # Fallback на простую генерацию
            fd, path = tempfile.mkstemp(suffix=".docx")
            os.close(fd)
            generate_test_docx(path, pages=1, text="Image Test")
            return path

        doc = Document()
        doc.add_heading("DOCX with Images", 0)

        # Создаем тестовое изображение
        img = Image.new('RGB', (100, 100), color='red')
        img_bytes = io.BytesIO()
        img.save(img_bytes, format='PNG')
        img_bytes.seek(0)

        doc.add_picture(img_bytes, width=Inches(1))

        fd, path = tempfile.mkstemp(suffix=".docx")
        os.close(fd)
        doc.save(path)
        yield path
        os.unlink(path)

    @pytest.fixture
    def large_docx(self) -> str:
        """Создание большого DOCX файла (>50 страниц)."""
        fd, path = tempfile.mkstemp(suffix=".docx")
        os.close(fd)
        generate_test_docx(path, pages=50, text="Large DOCX Document")
        yield path
        os.unlink(path)

    @pytest.mark.asyncio
    async def test_print_simple_docx(
        self,
        printer: FastPrinter,
        sample_docx: str,
        printer_ip: Optional[str],
    ) -> None:
        """Тест печати простого DOCX (текст)."""
        if not printer_ip:
            pytest.skip("Printer IP not specified")

        job = await printer.print_docx_fast(sample_docx, copies=1, duplex=False)

        assert job.status == PrintStatus.COMPLETED, f"Print failed: {job.error}"
        assert job.duration_ms is not None

    @pytest.mark.asyncio
    async def test_print_docx_with_tables(
        self,
        printer: FastPrinter,
        docx_with_tables: str,
        printer_ip: Optional[str],
    ) -> None:
        """Тест печати DOCX с таблицами."""
        if not printer_ip:
            pytest.skip("Printer IP not specified")

        job = await printer.print_docx_fast(docx_with_tables, copies=1, duplex=False)

        assert job.status == PrintStatus.COMPLETED, f"Print failed: {job.error}"

    @pytest.mark.asyncio
    async def test_print_docx_with_images(
        self,
        printer: FastPrinter,
        docx_with_images: str,
        printer_ip: Optional[str],
    ) -> None:
        """Тест печати DOCX с изображениями."""
        if not printer_ip:
            pytest.skip("Printer IP not specified")

        job = await printer.print_docx_fast(docx_with_images, copies=1, duplex=False)

        assert job.status == PrintStatus.COMPLETED, f"Print failed: {job.error}"

    @pytest.mark.asyncio
    async def test_print_large_docx(
        self,
        printer: FastPrinter,
        large_docx: str,
        printer_ip: Optional[str],
    ) -> None:
        """Тест печати большого DOCX (>50 страниц)."""
        if not printer_ip:
            pytest.skip("Printer IP not specified")

        job = await printer.print_docx_fast(large_docx, copies=1, duplex=False)

        # Большой документ может занять больше времени
        assert job.status == PrintStatus.COMPLETED, f"Print failed: {job.error}"
        # 50 страниц должны напечататься за 2 минуты
        if job.duration_ms:
            assert job.duration_ms < 120000, f"Print took too long: {job.duration_ms}ms"

    @pytest.mark.asyncio
    async def test_docx_to_pdf_conversion(
        self,
        sample_docx: str,
    ) -> None:
        """Тест конвертации DOCX в PDF."""
        from core.converters import docx_to_pdf

        try:
            pdf_path = await docx_to_pdf(sample_docx)
            assert os.path.exists(pdf_path)
            assert pdf_path.endswith(".pdf")
            os.unlink(pdf_path)
        except RuntimeError as e:
            # LibreOffice может быть не установлен
            pytest.skip(f"Conversion not available: {e}")

    @pytest.mark.asyncio
    async def test_print_docx_multiple_copies(
        self,
        printer: FastPrinter,
        sample_docx: str,
        printer_ip: Optional[str],
    ) -> None:
        """Тест печати нескольких копий DOCX."""
        if not printer_ip:
            pytest.skip("Printer IP not specified")

        job = await printer.print_docx_fast(sample_docx, copies=3, duplex=False)

        assert job.status == PrintStatus.COMPLETED, f"Print failed: {job.error}"

    @pytest.mark.asyncio
    async def test_print_docx_with_duplex(
        self,
        printer: FastPrinter,
        sample_docx: str,
        printer_ip: Optional[str],
    ) -> None:
        """Тест печати DOCX с дуплексом."""
        if not printer_ip:
            pytest.skip("Printer IP not specified")

        job = await printer.print_docx_fast(sample_docx, copies=1, duplex=True)

        assert job.status == PrintStatus.COMPLETED, f"Print failed: {job.error}"
